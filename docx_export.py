"""Génération de rapports Word (.docx) — template IDENTIQUE au PDF Ramery.

Structure reproduite :
  - Header bleu pleine largeur : logo gauche + "Meet Notes - Ramery" droite
  - Footer : ligne rouge + date + numéro de page centré
  - Titre  : grand, bleu, centré + ligne rouge + date
  - H1     : bleu gras 14pt + filet rouge dessous
  - H2     : barre bleue gauche + fond bleu clair
  - H3     : sombre gras 10pt
  - Puces  : texte sombre indenté (• rouge simulé)
  - Tableau: en-tête bleu/blanc, lignes alternées blanc/bleu clair
  - ---    : séparateur gris clair
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips

# ── Palette (identique pdf_export.py) ─────────────────────────────────────────
BLUE  = RGBColor(25,  60, 108)
RED   = RGBColor(211, 36,  34)
WHITE = RGBColor(255, 255, 255)
DARK  = RGBColor(20,  30,  48)
MUTED = RGBColor(100, 120, 150)
LIGHT = RGBColor(235, 241, 250)

_HEX = {
    "blue":  "193C6C",
    "red":   "D32422",
    "white": "FFFFFF",
    "dark":  "141E30",
    "muted": "647896",
    "light": "EBF1FA",
    "lgray": "E8EDF3",
}

ASSETS_DIR = Path(__file__).parent / "assets"
LOGO_PATH  = next(
    (ASSETS_DIR / f for f in ["logo.png", "logo.jpg", "logo.jpeg"]
     if (ASSETS_DIR / f).exists()),
    None,
)


# ── Helpers XML ───────────────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _para_bg(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _para_border(para, side: str, hex_color: str, sz: str = "6", space: str = "1") -> None:
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), hex_color)
    pBdr.append(el)


def _no_table_borders(tbl) -> None:
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _page_number_field(para) -> None:
    run = para.add_run()
    for tag, attr in [
        ("w:fldChar",   {"w:fldCharType": "begin"}),
        ("w:instrText", None),
        ("w:fldChar",   {"w:fldCharType": "end"}),
    ]:
        el = OxmlElement(tag)
        if attr:
            for k, v in attr.items():
                el.set(qn(k), v)
        if tag == "w:instrText":
            el.text = " PAGE "
        run._r.append(el)


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
    text = re.sub(r"`(.+?)`",              r"\1", text)
    text = re.sub(r"_{1,2}(.+?)_{1,2}",   r"\1", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text.strip()


def _add_inline_bold(para, text: str, size_pt: int = 10, color: RGBColor = None) -> None:
    """Ajoute du texte avec gras inline (**bold**) dans un paragraphe DOCX."""
    if color is None:
        color = DARK
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        clean = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', part)
        if clean.startswith('**') and clean.endswith('**'):
            run = para.add_run(clean[2:-2])
            run.bold = True
        else:
            # Préserver les espaces, juste retirer les marqueurs markdown
            clean = re.sub(r'\*{1,2}(.+?)\*{1,2}', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            run = para.add_run(clean)
            run.bold = False
        run.font.size      = Pt(size_pt)
        run.font.color.rgb = color


# ── Header ────────────────────────────────────────────────────────────────────

def _build_header(doc: Document) -> None:
    """Header bleu pleine largeur : logo à gauche + texte à droite (= PDF)."""
    section = doc.sections[0]
    header  = section.header
    section.header_distance = Cm(0)

    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Table 1 ligne × 2 colonnes dans le header (logo | texte)
    tbl = header.add_table(rows=1, cols=2, width=Inches(7.7))
    _no_table_borders(tbl)

    # Largeurs : 1.5 cm logo | reste texte
    tbl.columns[0].width = Cm(1.8)
    tbl.columns[1].width = Cm(15.0)

    cell_logo = tbl.cell(0, 0)
    cell_txt  = tbl.cell(0, 1)

    _cell_bg(cell_logo, _HEX["blue"])
    _cell_bg(cell_txt,  _HEX["blue"])

    # Logo
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(1)
    p_logo.paragraph_format.space_after  = Pt(1)
    _para_bg(p_logo, _HEX["blue"])
    if LOGO_PATH:
        try:
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(LOGO_PATH), height=Cm(0.85))
        except Exception:
            pass

    # Texte "Meet Notes  -  Ramery"
    p_txt = cell_txt.paragraphs[0]
    p_txt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_txt.paragraph_format.space_before = Pt(2)
    p_txt.paragraph_format.space_after  = Pt(2)
    _para_bg(p_txt, _HEX["blue"])
    run = p_txt.add_run("Meet Notes  -  Ramery")
    run.bold           = True
    run.font.size      = Pt(8)
    run.font.color.rgb = WHITE


# ── Footer ────────────────────────────────────────────────────────────────────

def _build_footer(doc: Document) -> None:
    """Footer : ligne rouge + 'Généré le date · Page X' centré (= PDF)."""
    section = doc.sections[0]
    footer  = section.footer

    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)

    _para_border(p, "top", _HEX["red"], sz="8", space="4")

    date_str = datetime.now().strftime("%d/%m/%Y à %H:%M")
    r1 = p.add_run(f"Généré le {date_str}  ·  Page ")
    r1.font.size      = Pt(7)
    r1.font.color.rgb = MUTED

    _page_number_field(p)


# ── Tableau ───────────────────────────────────────────────────────────────────

def _render_table(doc: Document, md_lines: list[str]) -> None:
    rows_data: list[list[str]] = []
    for line in md_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(set(c.replace(":", "").replace("-", "")) == set() for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    n_cols = max(len(r) for r in rows_data)
    tbl    = doc.add_table(rows=len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"

    for r_i, row_data in enumerate(rows_data):
        row = tbl.rows[r_i]
        for c_i in range(n_cols):
            cell      = row.cells[c_i]
            cell_text = row_data[c_i] if c_i < len(row_data) else ""
            para      = cell.paragraphs[0]
            para.clear()
            run = para.add_run(_strip_md_inline(cell_text))

            if r_i == 0:
                _cell_bg(cell, _HEX["blue"])
                run.bold           = True
                run.font.size      = Pt(8)
                run.font.color.rgb = WHITE
            else:
                bg = _HEX["light"] if r_i % 2 == 0 else _HEX["white"]
                _cell_bg(cell, bg)
                run.font.size      = Pt(8)
                run.font.color.rgb = DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Rendu Markdown ────────────────────────────────────────────────────────────

def _render_markdown(doc: Document, markdown: str) -> None:
    table_buf: list[str] = []

    def flush_table():
        if table_buf:
            _render_table(doc, list(table_buf))
            table_buf.clear()

    for raw in markdown.splitlines():
        line = raw.strip()

        if line.startswith("|") and line.endswith("|"):
            table_buf.append(line)
            continue
        else:
            flush_table()

        try:
            # H1 — bleu gras + filet rouge dessous
            if line.startswith("# ") and not line.startswith("## "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(_strip_md_inline(line[2:]))
                run.bold           = True
                run.font.size      = Pt(14)
                run.font.color.rgb = BLUE
                _para_border(p, "bottom", _HEX["red"], sz="6")

            # H2 — barre bleue gauche + fond bleu clair
            elif line.startswith("## "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(4)
                _para_bg(p, _HEX["light"])
                _para_border(p, "left", _HEX["blue"], sz="24", space="4")
                run = p.add_run(_strip_md_inline(line[3:]))
                run.bold           = True
                run.font.size      = Pt(11)
                run.font.color.rgb = BLUE

            # H3 — sombre gras
            elif line.startswith("### "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                run = p.add_run(_strip_md_inline(line[4:]))
                run.bold           = True
                run.font.size      = Pt(10)
                run.font.color.rgb = DARK

            # Sous-bullet (indentation dans raw, avant le bullet principal)
            elif re.match(r'^\s+[-*+]\s+', raw):
                text = re.sub(r'^\s+[-*+]\s+', '', raw)
                p    = doc.add_paragraph(style="List Bullet 2")
                p.paragraph_format.space_after = Pt(1)
                _add_inline_bold(p, text, size_pt=9, color=DARK)

            # Bullet principal
            elif re.match(r'^[-*+]\s+', raw):
                text = re.sub(r'^[-*+]\s+', '', raw)
                p    = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                _add_inline_bold(p, text, size_pt=10, color=DARK)

            # Ligne entièrement en gras → même style H2 bleu
            elif re.match(r'^\*\*[^*]+\*\*\s*:?\s*$', line):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(4)
                _para_bg(p, _HEX["light"])
                _para_border(p, "left", _HEX["blue"], sz="24", space="4")
                run = p.add_run(_strip_md_inline(line))
                run.bold           = True
                run.font.size      = Pt(11)
                run.font.color.rgb = BLUE

            # Séparateur ---
            elif line.startswith("---"):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(4)
                _para_border(p, "bottom", _HEX["lgray"], sz="4")

            # Ligne vide
            elif line == "":
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)

            # Texte avec gras inline
            elif '**' in line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                _add_inline_bold(p, line, size_pt=10, color=DARK)

            # Texte normal
            else:
                p   = doc.add_paragraph()
                run = p.add_run(_strip_md_inline(line))
                run.font.size      = Pt(10)
                run.font.color.rgb = DARK

        except Exception as e:
            print(f"[DOCX] Ligne ignorée : {line[:60]} — {e}")
            continue

    flush_table()


# ── Point d'entrée ────────────────────────────────────────────────────────────

def generate_docx(report_md: str, title: str) -> bytes:
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _build_header(doc)
    _build_footer(doc)

    # ── Page de titre ─────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title.replace("_", " "))
    run.bold           = True
    run.font.size      = Pt(20)
    run.font.color.rgb = BLUE

    # Ligne rouge sous le titre
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(2)
    sep.paragraph_format.space_after  = Pt(6)
    _para_border(sep, "bottom", _HEX["red"], sz="10")

    # Date
    date_str = datetime.now().strftime("%d/%m/%Y à %H:%M")
    p_date   = doc.add_paragraph(f"Compte rendu généré le {date_str}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_date.runs:
        p_date.runs[0].font.size      = Pt(9)
        p_date.runs[0].font.color.rgb = MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── Corps du rapport ──────────────────────────────────────────────────────
    _render_markdown(doc, report_md)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
